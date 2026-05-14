from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ───────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Style helpers ──────────────────────────────────────────────────────────
def set_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, header_color='1F3864', alt_color='DCE6F1'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], header_color)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            if ri % 2 == 1:
                shade_cell(cells[ci], alt_color)
    return table

# ════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROGRAMA DE ESTUDIO")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 56, 100)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Diplomado IA aplicada a la Gestión de Personas — UC")
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(31, 56, 100)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("EXAMEN: 14 de mayo de 2026")
run3.bold = True
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("Profesor: Raicho Bojilov")
run4.font.size = Pt(12)
run4.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run("Preparado con ayuda de Claude Code — mayo 2026")
run5.italic = True
run5.font.size = Pt(10)
run5.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  ÍNDICE
# ════════════════════════════════════════════════════════════════════════════
set_heading(doc, "ÍNDICE DEL PROGRAMA", level=1, color=(31,56,100))
index_items = [
    "Visión General del Programa (4 días)",
    "DÍA 1 — Sábado 10 mayo: Clases 1 y 2",
    "DÍA 2 — Domingo 11 mayo: Clases 3 y 4 (Caso Testscore)",
    "DÍA 3 — Lunes 12 mayo: Clase 5 (Programas de referidos y entrevistas)",
    "DÍA 4 — Martes 13 mayo: Clase 6 + Repaso general (Incentivos y discriminación)",
    "Tablas de Respuestas de todos los Controles",
    "Errores frecuentes a evitar",
]
for item in index_items:
    add_bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  VISIÓN GENERAL
# ════════════════════════════════════════════════════════════════════════════
set_heading(doc, "VISIÓN GENERAL DEL PROGRAMA", level=1, color=(31,56,100))
add_paragraph(doc,
    "El examen cubre 6 clases completas. El programa distribuye el estudio en 4 días, "
    "repasando dos clases por día los primeros tres días y dejando el último día para "
    "la clase más difícil (Clase 6) y un repaso final de todos los controles.",
    size=11)
doc.add_paragraph()

headers_gen = ["Día", "Fecha", "Clases", "Temas principales"]
rows_gen = [
    ["Día 1", "Sáb 10/5", "1 y 2", "¿Qué es la IA? Usos/abusos de datos y causalidad"],
    ["Día 2", "Dom 11/5", "3 y 4", "Análisis descriptivo + regresiones (Caso Testscore)"],
    ["Día 3", "Lun 12/5", "5",     "Mercado laboral, referidos, entrevistas vs IA"],
    ["Día 4", "Mar 13/5", "6 + repaso", "Incentivos, Safelite, discriminación + repaso total"],
]
add_table(doc, headers_gen, rows_gen)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  DÍA 1
# ════════════════════════════════════════════════════════════════════════════
set_heading(doc, "DÍA 1 — SÁBADO 10 DE MAYO", level=1, color=(31,56,100))
set_heading(doc, "Clases 1 y 2: ¿Qué es la IA en RRHH? Usos y Abusos de los Datos", level=2, color=(0,70,127))

# --- CLASE 1 ---
set_heading(doc, "CLASE 1: La IA en la Gestión de Personas", level=3)

add_paragraph(doc, "CONCEPTO 1 — ¿Qué hace la IA en Recursos Humanos?", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Imagina que tienes 1.000 postulantes a un cargo. No puedes entrevistar a todos. "
    "La IA analiza datos de los candidatos (notas, experiencia, tests) y te dice cuáles "
    "tienen más probabilidades de ser buenos empleados. Hace lo mismo que haría un reclutador "
    "experto, pero más rápido y con más datos.", size=11)
add_bullet(doc, "Predicción: ¿quién tendrá buen desempeño? ¿quién se quedará más tiempo?", "¿Qué predice? ")
add_bullet(doc, "No solo tecnología — es un proceso que incluye datos, modelos y decisiones humanas.", "Importante: ")
add_bullet(doc,
    "Reconocer patrones en datos de empleados pasados → aplicar esos patrones a candidatos nuevos.",
    "Cómo funciona: ")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 2 — El algoritmo Random Forest (Bosque Aleatorio)", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Imagina que preguntas a 500 amigos: '¿Crees que Juan será buen empleado?'. "
    "Cada amigo tiene información distinta. Al final, gana la opinión de la mayoría. "
    "Eso es el Random Forest: muchos 'árboles de decisión' que votan por la respuesta.", size=11)
add_bullet(doc, "No depende de una sola regla → es más robusto que un solo modelo.", "Ventaja: ")
add_bullet(doc, "No hay una fórmula simple que explique la decisión.", "Desventaja: ")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 3 — Las 4 etapas de un proyecto analítico", bold=True, color=(192,0,0))
add_table(doc,
    ["Etapa", "¿Qué se hace?", "Ejemplo Testscore"],
    [
        ["1. Análisis preliminar", "Explorar si hay datos útiles y si vale la pena el proyecto",
         "¿Tenemos datos de candidatos anteriores con resultados conocidos?"],
        ["2. Desarrollo", "Construir el modelo o herramienta",
         "Construir el algoritmo Testscore 1-100"],
        ["3. Implementación", "Usar la herramienta en decisiones reales",
         "Aplicar Testscore a nuevos candidatos"],
        ["4. Evaluación", "Medir si la herramienta realmente funciona",
         "¿El Testscore predice retención? (Spoiler: poco)"],
    ]
)

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 4 — ¿Cuándo conviene usar IA? (Regla de los 3 ingredientes)", bold=True, color=(192,0,0))
add_bullet(doc, "¿Tenemos datos históricos suficientes?", "1. Datos: ")
add_bullet(doc, "¿El problema es suficientemente complejo para que la IA ayude?", "2. Complejidad: ")
add_bullet(doc, "¿Los beneficios superan los costos (implementación, sesgos, errores)?", "3. Valor: ")

doc.add_paragraph()
add_paragraph(doc, "Preguntas del Control 1 (para repasar la lógica):", bold=True, italic=True)
add_bullet(doc, "¿Qué predice la IA? → rendimiento y/o permanencia en el trabajo.")
add_bullet(doc, "¿Qué es Random Forest? → conjunto de árboles de decisión que 'votan'.")
add_bullet(doc, "¿Para qué sirven las 4 etapas? → para ordenar el proceso analítico.")

doc.add_paragraph()

# --- CLASE 2 ---
set_heading(doc, "CLASE 2: Usos y Abusos de los Datos y la Causalidad", level=3)

add_paragraph(doc, "CONCEPTO 5 — Correlación ≠ Causalidad (el error más común)", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Ejemplo clásico: las personas que van al hospital mueren más que las que no van. "
    "Conclusión equivocada: 'el hospital mata gente'. Conclusión correcta: van al hospital "
    "porque están enfermos. La causalidad requiere más evidencia que una correlación.", size=11)
add_bullet(doc, "Correlación: dos cosas se mueven juntas (pueden ser independientes).")
add_bullet(doc, "Causalidad: una cosa CAUSA la otra (mucho más difícil de probar).")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 6 — Diferencias en Diferencias (DiD)", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Imagina que quieres saber si un entrenamiento mejoró el desempeño. "
    "Comparas el desempeño antes y después, PERO también comparas con un grupo que NO recibió "
    "el entrenamiento. La 'diferencia en diferencias' elimina otros factores que habrían "
    "cambiado igual (como que llegó temporada alta).", size=11)
add_bullet(doc, "Fórmula: (Grupo A después - Grupo A antes) - (Grupo B después - Grupo B antes)")
add_bullet(doc, "Supuesto clave: sin el tratamiento, ambos grupos habrían evolucionado igual ('tendencias paralelas').")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 7 — Regresión Discontinua (RD)", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Si hay un corte arbitrario (ej. puntaje 600 para entrar a la universidad), "
    "los que sacaron 598 vs 601 son casi idénticos, excepto que uno entró y el otro no. "
    "Comparar esos dos grupos aislado del 'efecto de estar cerca del corte'.", size=11)
add_bullet(doc, "Requiere que el corte sea aleatorio o arbitrario, no manipulable.")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 8 — Variable Instrumental (IV)", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Necesitas una 'palanca' externa que mueve tu variable de interés pero que no tiene "
    "ningún otro efecto directo. Ejemplo: distancia a la universidad instrumentaliza "
    "el nivel educacional (quien vive lejos estudia menos, pero la distancia en sí no "
    "afecta el salario directamente).", size=11)

doc.add_paragraph()
add_paragraph(doc, "Preguntas del Control 2 (lógica clave):", bold=True, italic=True)
add_bullet(doc, "¿Por qué no basta la correlación? → no prueba causalidad.")
add_bullet(doc, "¿Cuándo usar DiD? → cuando hay grupos comparables y tratamiento en un subconjunto.")
add_bullet(doc, "¿Cuándo usar IV? → cuando hay selección en el tratamiento y se tiene una palanca exógena.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  DÍA 2
# ════════════════════════════════════════════════════════════════════════════
set_heading(doc, "DÍA 2 — DOMINGO 11 DE MAYO", level=1, color=(31,56,100))
set_heading(doc, "Clases 3 y 4: Análisis Descriptivo y Regresiones (Caso Testscore)", level=2, color=(0,70,127))

# --- CLASE 3 ---
set_heading(doc, "CLASE 3: Análisis Descriptivo — El primer paso de cualquier análisis", level=3)

add_paragraph(doc,
    "CONCEPTO 9 — ¿Para qué sirve el análisis descriptivo? (FUNDAMENTAL para el examen)",
    bold=True, color=(192,0,0))
add_paragraph(doc,
    "Antes de hacer modelos complicados, primero miras los datos como son. "
    "Si en esta etapa no encuentras ninguna relación interesante, casi seguro que no existe "
    "aunque uses modelos más avanzados.", size=11)
p = doc.add_paragraph()
run = p.add_run("Cita directa de la Clase 3 (p.9): ")
run.bold = True
p.add_run('"Si no se encuentra ninguna relación significativa y relevante en los datos, '
          'es muy probable que no exista, incluso si se utilizan herramientas más avanzadas."')
add_bullet(doc, "Análisis básico: media, desviación estándar, mínimo, máximo. NO incluye correlaciones (eso viene después).")
add_bullet(doc, "El análisis descriptivo se usa en TODAS las etapas del proyecto, no solo al inicio.")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 10 — Histogramas vs Gráficos de Densidad", bold=True, color=(192,0,0))
add_table(doc,
    ["Característica", "Histograma", "Gráfico de Densidad"],
    [
        ["¿Qué muestra?", "Conteos brutos por intervalo (bins)", "Frecuencia relativa en curva suavizada"],
        ["Eje vertical", "Número de observaciones en cada bin", "Densidad (frecuencia relativa)"],
        ["Ventaja", "Fácil de entender", "Más fluido, facilita comparar distribuciones"],
        ["Para comparar grupos", "Difícil si hay grupos de distinto tamaño", "Ideal — normaliza por el total"],
    ]
)
doc.add_paragraph()
add_paragraph(doc,
    "Pregunta Control 3 Q5 (afirmación incorrecta de los gráficos de densidad): "
    "La respuesta es A — los gráficos de densidad NO muestran conteos brutos, eso lo hacen los histogramas.",
    italic=True, color=(89,89,89))

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 11 — Los Deciles", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Decil = dividir a TODAS las personas en 10 grupos iguales según su puntaje. "
    "El primer decil = el 10% con puntaje más bajo. El décimo decil = el 10% con puntaje más alto. "
    "Importante: se dividen las PERSONAS en grupos iguales, NO el rango de puntajes en intervalos iguales.",
    size=11)
add_paragraph(doc,
    "Pregunta Control 3 Q10 (afirmación incorrecta): La respuesta es B — los deciles dividen "
    "observaciones en grupos iguales, no el rango máximo-mínimo.",
    italic=True, color=(89,89,89))

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 12 — Tasas del Testscore (Control 3 Q2 y Q6)", bold=True, color=(192,0,0))
add_table(doc,
    ["Tasa", "Fórmula", "Tipo"],
    [
        ["Tasa de oferta incondicional", "Candidatos con oferta / Total candidatos", "Incondicional (sobre el total)"],
        ["Tasa de retención condicional", "Candidatos retenidos / Candidatos contratados", "Condicional (sobre etapa anterior)"],
        ["Tasa de retención incondicional", "Candidatos retenidos / Total candidatos", "Incondicional (sobre el total)"],
    ]
)

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 13 — Resultados reales del Testscore (¡importante!)", bold=True, color=(192,0,0))
add_bullet(doc, "Correlación Testscore-Retención: solo 0.15 (muy débil)")
add_bullet(doc, "Correlación Testscore-Desempeño: prácticamente 0 (no predice desempeño)")
add_bullet(doc, "Conclusión: el Testscore sirve un poco para filtrar en la etapa de entrevistas, pero NO para predecir quién va a rendir bien o quedarse mucho tiempo.")

doc.add_paragraph()
add_paragraph(doc,
    "Control 3 Q1 — Si los reclutadores siguen la política, ¿qué pasa con los candidatos de bajo Testscore? "
    "Respuesta E: su frecuencia entre entrevistados Y contratados debería ser CERO o cercana a cero "
    "(porque el corte los elimina en la primera etapa).",
    italic=True, color=(89,89,89))

doc.add_paragraph()

# --- CLASE 4 ---
set_heading(doc, "CLASE 4: Regresión Lineal — La herramienta principal del análisis cuantitativo", level=3)

add_paragraph(doc, "CONCEPTO 14 — ¿Qué es una regresión?", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Una regresión mide cuánto cambia Y cuando X aumenta en 1 unidad, manteniendo lo demás fijo. "
    "Ejemplo: por cada 10 puntos más de Testscore, ¿cuántos días más se queda el empleado? "
    "La respuesta viene del coeficiente de regresión.", size=11)

add_table(doc,
    ["Concepto", "¿Qué significa?", "Valor 'bueno'"],
    [
        ["Coeficiente (β)", "Cuánto sube Y cuando X sube 1 unidad", "Depende del contexto"],
        ["t-estadístico", "¿Es el coeficiente estadísticamente diferente de cero?", "Mayor a 1.96 (en valor absoluto)"],
        ["R²", "¿Qué % de la variación de Y explica el modelo?", "Depende; 0.15 es muy bajo"],
        ["p-valor", "Probabilidad de obtener ese resultado por azar", "Menor a 0.05 → significativo"],
        ["Intervalo de confianza", "Rango donde probablemente está el verdadero coeficiente", "No debe cruzar el cero"],
    ]
)

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 15 — El protocolo de 8 pasos para evaluar herramientas analíticas", bold=True, color=(192,0,0))
steps = [
    ("Paso 1:", "Definir el objetivo — ¿qué quieres predecir?"),
    ("Paso 2:", "Describir los datos disponibles"),
    ("Paso 3:", "Análisis descriptivo (histogramas, correlaciones)"),
    ("Paso 4:", "Correlaciones entre la herramienta y el outcome"),
    ("Paso 5:", "Regresión simple — herramienta vs outcome"),
    ("Paso 6:", "Regresión múltiple — controlando por otras variables"),
    ("Paso 7:", "Análisis por subgrupos (deciles)"),
    ("Paso 8:", "Conclusión y recomendación"),
]
for step, desc in steps:
    add_bullet(doc, desc, step + " ")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 16 — Regresión Multivariada: controlar por otras variables", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Si quieres saber si el Testscore predice retención, pero hay candidatos más jóvenes o con más "
    "experiencia en el grupo de alto Testscore, la regresión multivariada 'mantiene fija' la edad "
    "y la experiencia para aislar el efecto real del Testscore.", size=11)

doc.add_paragraph()
add_paragraph(doc,
    "Resultado clave Clase 4: El Testscore NO predice ni retención ni desempeño de forma robusta. "
    "El R² es muy bajo y los coeficientes no son estadísticamente significativos en la mayoría de especificaciones.",
    bold=True, italic=True, color=(192,0,0))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  DÍA 3
# ════════════════════════════════════════════════════════════════════════════
set_heading(doc, "DÍA 3 — LUNES 12 DE MAYO", level=1, color=(31,56,100))
set_heading(doc, "Clase 5: Mercado Laboral, Programas de Referidos y Entrevistas", level=2, color=(0,70,127))

add_paragraph(doc, "CONCEPTO 17 — El problema de información en el mercado laboral", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Cuando una empresa quiere contratar, tiene un problema: no sabe si el candidato "
    "realmente es bueno. El candidato sí lo sabe. Esto se llama asimetría de información "
    "(una parte sabe más que la otra). La empresa usa señales para reducir esta incertidumbre.",
    size=11)
add_bullet(doc, "Señal costosa: algo que solo los candidatos buenos pueden hacer (ej. graduarse de una universidad difícil). El diploma 'señaliza' calidad porque es costoso obtenerlo.")
add_bullet(doc, "Señal barata: algo que cualquiera puede imitar fácilmente → no tiene valor como señal.")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 18 — Las señales que usan las empresas", bold=True, color=(192,0,0))
add_table(doc,
    ["Señal", "¿Cómo funciona?", "¿Es costosa?"],
    [
        ["Educación (título)", "Solo los capaces pueden completarla", "Sí (tiempo, dinero, esfuerzo)"],
        ["Historia laboral", "Muestra comportamiento real en trabajos anteriores", "Indirectamente"],
        ["Referencias", "Un empleado actual avala al candidato", "Sí (reputación del referente)"],
        ["Test de aptitud", "Puntaje en prueba estandarizada", "Moderadamente"],
    ]
)

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 19 — Programa de Referidos (Estudio: Burks et al. 2015)", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Estudio en 3 industrias: call centers, transporte de carga y TI. "
    "¿Qué pasa cuando un empleado actual recomienda a alguien?", size=11)
add_bullet(doc, "Más bajo (más selectivos en a quién recomiendan)", "Tasa de contratación: ")
add_bullet(doc, "Más alta (si te van a echar al referido, tu reputación se daña)", "Tasa de aceptación de oferta: ")
add_bullet(doc, "MAYOR que los no referidos (el referente tiene info privada del candidato)", "Retención: ")
add_bullet(doc, "IGUAL o solo levemente mejor", "Desempeño: ")

p = doc.add_paragraph()
run = p.add_run("Clave del examen: ")
run.bold = True
p.add_run("Los referidos se quedan más tiempo, pero no necesariamente rinden mejor. "
          "La ventaja es en RETENCIÓN, no en desempeño.")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 20 — Tasa de excepción en entrevistas (Estudio: Hoffman, Kahn, Li 2018)", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Los algoritmos recomiendan a ciertos candidatos, pero los reclutadores a veces "
    "hacen 'excepciones' (contratan a alguien que el algoritmo no recomienda). "
    "¿Qué pasa con esos candidatos?", size=11)
add_bullet(doc, "Los empleados contratados por excepción (ignorando al algoritmo) tienen PEOR desempeño y mayor rotación.", "Resultado: ")
add_bullet(doc, "Cuando los gerentes tienen más autonomía para hacer excepciones, hacen MÁS excepciones.", "Hallazgo adicional: ")
add_bullet(doc,
    "Cita Clase 5: 'La tasa de excepción es mayor cuando los gerentes tienen más autonomía en la decisión de contratación.'",
    "Cita directa: ")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 21 — ¿Entrevistas vs IA?", bold=True, color=(192,0,0))
add_paragraph(doc,
    "El estudio de Clase 5 muestra que las entrevistas humanas no siempre superan al algoritmo. "
    "Los reclutadores confían en su 'intuición', pero esa intuición a menudo refleja sesgos "
    "o información irrelevante. Los algoritmos bien calibrados tienden a ser más consistentes.",
    size=11)
add_bullet(doc, "Entrevistas humanas: captura información cualitativa, pero introduce sesgos.")
add_bullet(doc, "Algoritmo: consistente, pero no capta todo el contexto.")
add_bullet(doc, "Mejor práctica: combinar ambos, pero respetar al algoritmo en decisiones estandarizadas.")

doc.add_paragraph()
add_paragraph(doc, "Preguntas clave de Control 5:", bold=True, italic=True)
add_bullet(doc, "¿Qué efectos tienen los programas de referidos? → Más retención, igual desempeño.")
add_bullet(doc, "¿Qué es una señal costosa? → Algo que solo los buenos candidatos pueden hacer.")
add_bullet(doc, "¿Qué pasa cuando la tasa de excepción es alta? → Peores resultados de contratación.")
add_bullet(doc, "¿Qué es la tasa de excepción? → Porcentaje de contrataciones que ignoran la recomendación del algoritmo.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  DÍA 4
# ════════════════════════════════════════════════════════════════════════════
set_heading(doc, "DÍA 4 — MARTES 13 DE MAYO", level=1, color=(31,56,100))
set_heading(doc, "Clase 6: Incentivos, Caso Safelite y Discriminación", level=2, color=(0,70,127))

# --- INCENTIVOS ---
set_heading(doc, "PARTE A: Teoría de Incentivos", level=3)

add_paragraph(doc, "CONCEPTO 22 — ¿Por qué importan los incentivos?", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Un empleado puede trabajar mucho o poco. Para que trabaje mucho, necesitas incentivos. "
    "Hay dos formas de pagar: salario fijo (sin incentivos) o pago variable (ligado al resultado). "
    "El pago variable aumenta el esfuerzo, pero también el riesgo para el empleado.",
    size=11)
add_bullet(doc, "Salario fijo: el empleado no tiene incentivo para esforzarse más.")
add_bullet(doc, "Pago por pieza (piece-rate): por cada unidad producida, recibes un monto.")
add_bullet(doc, "Bono por desempeño: si superas la meta, recibes un extra.")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 23 — Caso Safelite (Lazear 2000) — El más importante de Clase 6", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Safelite es una empresa que cambia parabrisas de autos. En 1994, cambió su sistema de pago "
    "de salario fijo a pago por pieza (por cada parabrisas instalado). ¿Qué pasó?",
    size=11)
add_table(doc,
    ["Efecto", "Resultado", "¿Por qué?"],
    [
        ["Productividad total", "+44% de aumento", "Efecto incentivo + efecto selección"],
        ["Efecto incentivo", "~24%", "Los mismos empleados trabajaron más"],
        ["Efecto selección", "~20%", "Los trabajadores menos productivos se fueron (no les convenía)"],
        ["Varianza del pago", "Aumentó significativamente", "Mayor dispersión en los ingresos"],
        ["Trabajadores más lentos", "Muchos se fueron", "El pago por pieza los penalizaba"],
    ]
)
doc.add_paragraph()
add_paragraph(doc,
    "Regla clave: pago variable = mayor esfuerzo + mayor selección. "
    "Los malos trabajadores se van solos porque no les rinde el sistema.",
    bold=True, color=(192,0,0))

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 24 — Incentivos Relativos: pagar según el grupo, no solo el individuo", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Si un día llueve, todos los vendedores venden menos, sin importar cuánto se esfuercen. "
    "Si el bono depende del resultado absoluto, los castigas injustamente. "
    "La solución: compararlos entre sí (torneos, rankings). El premio va al que vendió MÁS "
    "RELATIVO al grupo, no al que más vendió en términos absolutos.",
    size=11)
add_bullet(doc, "Elimina el efecto de shocks comunes (clima, economía, temporada).", "Ventaja: ")
add_bullet(doc, "Puede crear ambiente competitivo negativo (sabotaje, no compartir información).", "Desventaja: ")

doc.add_paragraph()

# --- DISCRIMINACIÓN ---
set_heading(doc, "PARTE B: Discriminación — Dos tipos muy distintos", level=3)

add_paragraph(doc, "CONCEPTO 25 — Discriminación por preferencia vs discriminación estadística", bold=True, color=(192,0,0))
add_table(doc,
    ["Tipo", "¿Qué es?", "Ejemplo", "¿Se elimina con información?"],
    [
        ["Discriminación por preferencia",
         "El empleador rechaza a alguien por prejuicio personal, sin razón económica",
         "No contrato mujeres porque no me gustan",
         "NO — es ideológica"],
        ["Discriminación estadística",
         "El empleador usa el grupo como señal de características individuales desconocidas",
         "No contrato a alguien porque su universidad es desconocida y no sé si es buena",
         "SÍ — con más información individual"],
    ]
)

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 26 — Estudio Bordón y Braga (2020): reputación universitaria y discriminación estadística", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Enviaron currículums ficticios a empleadores chilenos. "
    "Mismas características, pero variando la universidad de egreso. Resultados:",
    size=11)
add_bullet(doc, "Graduados de universidades de ALTA reputación → más llamadas a entrevistas.")
add_bullet(doc, "Graduados de universidades de BAJA reputación → muchas menos llamadas.")
add_bullet(doc, "Esto es discriminación estadística: el empleador usa la reputación de la universidad como señal de calidad del candidato, porque no tiene información directa de su desempeño.")
add_bullet(doc, "Bordón y Braga usan un experimento (envío de currículums) para aislar el efecto causal de la reputación universitaria.")

doc.add_paragraph()
add_paragraph(doc, "CONCEPTO 27 — Hoffman y Tadelis (2021): habilidades gerenciales y rotación", bold=True, color=(192,0,0))
add_paragraph(doc,
    "Estudiaron si los gerentes con mejores habilidades interpersonales reducen la rotación de su equipo. "
    "Usaron métricas de desempeño como señal de habilidad gerencial.",
    size=11)
p = doc.add_paragraph()
run = p.add_run("Resultado central: ")
run.bold = True
p.add_run("Los gerentes con mejores habilidades tienen equipos con MENOR rotación. "
          "El efecto es significativo y relevante (NO limitado, como dice la respuesta incorrecta del control).")
add_bullet(doc,
    "Control 6 Q incorrecta: 'Las habilidades gerenciales tienen un efecto LIMITADO sobre la tasa de rotación' → FALSO. "
    "El efecto es SIGNIFICATIVO.",
    "ERROR FRECUENTE: ")
add_bullet(doc,
    "Usaron MÉTRICAS DE DESEMPEÑO (no solo 'el desempeño' en abstracto) como señal de habilidad.",
    "Señal usada: ")

doc.add_paragraph()
add_paragraph(doc, "Preguntas clave de Control 6:", bold=True, italic=True)
add_bullet(doc, "¿Qué encontró Lazear (2000)? → Aumento del 44% en productividad (efecto incentivo + selección).")
add_bullet(doc, "¿Para qué sirven los incentivos relativos? → Para eliminar shocks comunes del mercado.")
add_bullet(doc, "¿Qué tipo de discriminación usa Bordón y Braga? → Estadística.")
add_bullet(doc, "¿Cuál es la señal en Hoffman y Tadelis? → Métricas de desempeño (no 'el desempeño' en abstracto).")
add_bullet(doc, "¿Cuál es el efecto de habilidades gerenciales sobre rotación? → Significativo y negativo (más habilidad → menos rotación).")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  TABLAS DE RESPUESTAS
# ════════════════════════════════════════════════════════════════════════════
set_heading(doc, "TABLAS DE RESPUESTAS — TODOS LOS CONTROLES", level=1, color=(31,56,100))

add_paragraph(doc,
    "Resumen de respuestas correctas para todos los controles del diplomado.",
    bold=True, size=11)
doc.add_paragraph()

# Control 3
set_heading(doc, "Control 3 — Clases 3 y 4 (Análisis descriptivo y Testscore)", level=2, color=(0,70,127))
add_table(doc,
    ["Pregunta", "Respuesta", "Tema"],
    [
        ["1", "e", "Densidad de entrevistados/contratados con bajo Testscore → frecuencia cero"],
        ["2", "b", "Tasa de oferta incondicional = candidatos con oferta / total candidatos"],
        ["3", "e", "Si no hay relación descriptiva, probablemente no existe"],
        ["4", "c", "Estadísticas básicas NO incluyen correlaciones"],
        ["5", "a", "Gráficos de densidad NO muestran conteos brutos (eso es el histograma)"],
        ["6", "c", "Tasa de retención condicional = retenidos / contratados"],
        ["7", "e", "Análisis descriptivo NO se usa SOLO en etapa preliminar"],
        ["8", "d", "Correlación mide relación lineal entre dos variables"],
        ["9", "c", "Eje vertical del histograma = número de observaciones por intervalo, NO frecuencia individual"],
        ["10", "b", "Deciles dividen observaciones en grupos iguales, NO el rango en intervalos iguales"],
    ]
)

doc.add_paragraph()

# Control 5
set_heading(doc, "Control 5 — Clase 5 (Referidos, señales, excepción)", level=2, color=(0,70,127))
add_table(doc,
    ["Pregunta", "Respuesta", "Tema"],
    [
        ["1", "c", "Programas de referidos aumentan la retención (no el desempeño de forma significativa)"],
        ["2", "b", "Las referencias son señales costosas (reputación del referente en juego)"],
        ["3", "d", "Tasa de excepción = contratados fuera de recomendación / total contratados"],
        ["4", "a", "Los referidos son más selectivos porque afectan la reputación del referente"],
        ["5", "e", "La info privada del referente explica la ventaja en retención de los referidos"],
        ["6", "b", "Mayor autonomía gerencial → mayor tasa de excepción"],
        ["7", "c", "La educación es una señal costosa (solo los capaces pueden completarla)"],
        ["8", "d", "Los contratados por excepción tienen peor desempeño que los recomendados por el algoritmo"],
        ["9", "a", "La asimetría de información favorece al candidato (él sabe más sobre sí mismo)"],
        ["10", "e", "El programa de referidos reduce los costos de búsqueda de la empresa"],
    ]
)
add_paragraph(doc, "* Las respuestas del Control 5 son aproximadas — verificar contra tus respuestas marcadas.", italic=True, color=(128,128,128))

doc.add_paragraph()

# Control 6
set_heading(doc, "Control 6 — Clase 6 (Incentivos, Safelite, discriminación)", level=2, color=(0,70,127))
add_table(doc,
    ["Preg. (v1)", "Preg. (v2 reorden)", "Respuesta", "Tema"],
    [
        ["1", "–", "b/c", "Definición efecto incentivo en Lazear"],
        ["2", "–", "a", "Definición efecto selección en Lazear"],
        ["3", "–", "d", "Aumento total de productividad en Safelite (+44%)"],
        ["4", "–", "c", "Para qué sirven incentivos relativos (shocks comunes)"],
        ["5", "–", "b", "Discriminación estadística usa el grupo como señal"],
        ["6", "–", "e/c", "Señal en Hoffman & Tadelis = métricas de desempeño"],
        ["7", "–", "a", "Bordón y Braga: experimento con currículums"],
        ["8", "–", "a (NO b)", "Hoffman & Tadelis: efecto de habilidades gerenciales es SIGNIFICATIVO (no limitado)"],
        ["9", "–", "d", "Safelite: efecto sobre varianza del pago"],
        ["10", "–", "b (NO e)", "Señal en Hoffman & Tadelis: MÉTRICAS de desempeño (no 'el desempeño')"],
    ]
)
add_paragraph(doc, "ERRORES FRECUENTES en Control 6:", bold=True, color=(192,0,0))
add_bullet(doc,
    "Q8: La respuesta INCORRECTA es b ('efecto limitado de habilidades gerenciales'). "
    "El efecto ES significativo según la clase. La afirmación incorrecta es A ('efecto limitado').",
    "⚠ ")
add_bullet(doc,
    "Q10: La respuesta es b ('métricas de desempeño') NO e ('el desempeño'). "
    "La distinción es precisa: en el contexto de sistemas de incentivos se usan métricas medibles.",
    "⚠ ")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  REPASO FINAL
# ════════════════════════════════════════════════════════════════════════════
set_heading(doc, "REPASO FINAL — CONCEPTOS MÁS IMPORTANTES POR CLASE", level=1, color=(31,56,100))

add_table(doc,
    ["Clase", "Concepto clave", "Pregunta típica de examen"],
    [
        ["1 – IA en RRHH",
         "4 etapas del proyecto; Random Forest; IA predice desempeño/retención",
         "¿En qué etapa se evalúa la herramienta?"],
        ["2 – Causalidad",
         "Correlación ≠ causalidad; DiD requiere tendencias paralelas; IV necesita instrumento exógeno",
         "¿Por qué no basta una correlación para probar causalidad?"],
        ["3 – Descriptivo",
         "Estadísticas básicas (no incluyen correlaciones); histogramas vs densidad; deciles = grupos iguales de personas",
         "¿Qué muestra el eje vertical de un histograma?"],
        ["4 – Regresión",
         "t-estadístico > 1.96 → significativo; R² bajo = poca predicción; protocolo 8 pasos",
         "¿Qué indica un R² de 0.02?"],
        ["5 – Referidos",
         "Referidos tienen mayor retención; excepción al algoritmo → peor desempeño; señales costosas",
         "¿Qué efecto tienen los programas de referidos sobre el desempeño?"],
        ["6 – Incentivos",
         "Safelite +44% (incentivo 24% + selección 20%); discriminación estadística se elimina con info; habilidades gerenciales → rotación",
         "¿Cuál fue el efecto total en Safelite y cómo se descompone?"],
    ]
)

doc.add_paragraph()
set_heading(doc, "Los 5 errores más comunes en los controles", level=2, color=(192,0,0))
errors = [
    ("Deciles", "Los deciles NO dividen el RANGO en partes iguales. Dividen a las PERSONAS en grupos iguales."),
    ("Histogramas", "Los gráficos de densidad NO muestran conteos brutos. Los histogramas SÍ."),
    ("Análisis descriptivo", "El análisis descriptivo NO se usa solo en la etapa preliminar — se usa en todas."),
    ("Correlación", "La correlación mide relación lineal. NO implica causalidad. NO equivale al coeficiente de regresión."),
    ("Señal Hoffman & Tadelis", "La señal son MÉTRICAS DE DESEMPEÑO (no 'el desempeño' en abstracto). Y el efecto es SIGNIFICATIVO, no limitado."),
]
for term, desc in errors:
    add_bullet(doc, desc, term + ": ")

doc.add_paragraph()
set_heading(doc, "Frase final para el examen", level=2, color=(0,70,127))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    '"Si el análisis descriptivo no revela ninguna relación, es muy probable que no exista, '
    'incluso usando herramientas más avanzadas." — Clase 3, p. 9'
)
run.bold = True
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(31, 56, 100)

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("¡Mucho éxito en el examen del 14 de mayo!")
run2.bold = True
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0, 176, 80)

# ─── Save ────────────────────────────────────────────────────────────────────
output_path = (r"c:\Users\amsasadykov\OneDrive - Grupo Minero Antofagasta Minerals"
               r"\Escritorio\Proyectos Claude\Cursos y capacitaciones\Diplomado UC"
               r"\Programa_Estudio_Examen_14Mayo.docx")
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
